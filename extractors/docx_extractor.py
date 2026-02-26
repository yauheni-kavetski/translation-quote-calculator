import io
from docx import Document

def extract_docx(file_bytes: bytes) -> str:
    """
    Извлекает текст и таблицы из docx.
    """
    doc = Document(io.BytesIO(file_bytes))
    content = "\n".join(p.text for p in doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text for cell in row.cells)
            if row_text.strip():
                content += row_text + "\n"
    return content